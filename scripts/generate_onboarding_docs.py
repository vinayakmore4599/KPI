#!/usr/bin/env python3
"""Generate onboarding Word documents from markdown sources in docs/onboarding/."""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
ONBOARDING_DIR = ROOT / "docs" / "onboarding"

sys.path.insert(0, str(ROOT / "scripts"))
from docx_helpers import add_code_block, add_table, add_title_page, configure_document  # noqa: E402

from docx import Document  # noqa: E402

SOURCES = [
    (
        "onboarding-column-function.md",
        "Onboarding a New Column Function",
        "Row-wise transforms before aggregation (base_measures.op)",
    ),
    (
        "onboarding-measure-function.md",
        "Onboarding a New Measure Function",
        "Scalar math over computed measures (op: fn / arithmetic / expr)",
    ),
    (
        "onboarding-hook.md",
        "Onboarding a New Hook",
        "Aligned period-series algorithms (op: hook)",
    ),
    (
        "onboarding-op.md",
        "Onboarding a New Op",
        "First-class measure kinds (OpPlugin lifecycle)",
    ),
]


def _split_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    inner = stripped.strip("|").replace(" ", "")
    return bool(inner) and set(inner) <= {"-", ":", "|"}


def markdown_to_docx(md_path: Path, doc: Document) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    i = 0
    skip_title = True
    in_code = False
    code_lines: list[str] = []
    table_headers: list[str] | None = None
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_headers, table_rows
        if table_headers is not None:
            add_table(doc, table_headers, table_rows)
            table_headers = None
            table_rows = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line.rstrip())
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            i += 1
            continue

        if line.strip() == "---":
            flush_table()
            i += 1
            continue

        if line.startswith("|") and _is_table_separator(line):
            i += 1
            continue

        if line.startswith("|"):
            flush_table()
            cells = _split_table_row(line)
            if table_headers is None:
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue

        if table_headers is not None and not line.startswith("|"):
            flush_table()

        if skip_title and line.startswith("# "):
            skip_title = False
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(line.strip(), style="List Number")
        elif line.startswith("- [ ] "):
            doc.add_paragraph(line[6:].strip(), style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip().startswith("**Audience:**") or line.strip().startswith("**Version:**"):
            p = doc.add_paragraph()
            p.add_run(line.strip()).italic = True
        elif line.strip():
            doc.add_paragraph(line.strip())
        i += 1

    flush_table()
    if in_code and code_lines:
        add_code_block(doc, code_lines)


def build_one(source_name: str, title: str, subtitle: str) -> Path:
    md_path = ONBOARDING_DIR / source_name
    if not md_path.exists():
        raise FileNotFoundError(md_path)

    doc = Document()
    configure_document(doc)
    add_title_page(doc, title, subtitle, version=date.today().strftime("%B %Y"))
    markdown_to_docx(md_path, doc)

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main() -> None:
    ONBOARDING_DIR.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []
    for source, title, subtitle in SOURCES:
        out = build_one(source, title, subtitle)
        outputs.append(out)
        print(f"Wrote {out.relative_to(ROOT)}")
    print(f"Generated {len(outputs)} onboarding documents.")


if __name__ == "__main__":
    main()
