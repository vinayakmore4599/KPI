#!/usr/bin/env python3
"""Generate KPI Engine design document (Word) for architecture and leadership."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CAPABILITIES = ROOT / "kpi_engine/registries/CAPABILITIES.md"
OUTPUT = ROOT / "docs/KPI-Engine-Design-Document.docx"


def _parse_capabilities(path: Path) -> dict[str, list[str]]:
    text = path.read_text(encoding="utf-8")
    sections: dict[str, list[str]] = {}
    current = ""
    for line in text.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif line.startswith("### `") and current:
            name = line.split("`")[1]
            sections[current].append(name)
    return sections


def _add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x44, 0x54, 0x6A)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(f"Version: August 2026  |  Audience: Leadership, Architecture, Data Engineering")
    m.font.size = Pt(10)
    m.italic = True
    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    caps = _parse_capabilities(CAPABILITIES)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    _add_title(
        doc,
        "KPI Engine",
        "Config-Driven KPI Calculation Framework\nDesign Principles, Capabilities & Onboarding",
    )

    # 1 Executive summary
    _heading(doc, "1. Executive Summary")
    _para(
        doc,
        "The KPI Engine is a reusable, config-driven calculation library that turns a standard "
        "request context and KPI YAML into JSON results for dashboards and reports. It separates "
        "what to calculate (YAML authored by data engineers) from how the platform runs it "
        "(DuckDB extract + Pandas calculation).",
    )
    _bullets(
        doc,
        [
            "One generic entry point (kpi_engine.main) serves every KPI; routing is by kpi_id.",
            "New KPIs are onboarded primarily by YAML — typically ~20 lines for a simple metric.",
            "702 automated tests validate bind rules, SQL generation, and calculation semantics.",
            "Horizontally scalable: each request is stateless; Kubernetes scales worker pods.",
            "Extensible via an allowlisted capability catalog — no per-KPI Python forks.",
        ],
    )

    # 2 Objectives
    _heading(doc, "2. Strategic Objectives")
    _table(
        doc,
        ["Objective", "How the engine delivers"],
        [
            ["Reusability", "Single engine + registries; KPI logic lives in YAML, not duplicated code."],
            ["Scalability", "Stateless compute(context) per request; scale K8s replicas; no shared request state."],
            ["Consistency", "One anchor month, one scan width, deterministic JSON contract for every KPI."],
            ["Governance", "Allowlisted ops/functions/hooks; bind-time validation; no eval or import paths from YAML."],
            ["Velocity", "Copy template KPI YAML, validate without ADLS, ship; ~20 lines for simple KPIs."],
            ["Separation of concerns", "Platform owns context, ADLS, auth; engine owns bind → extract → calculate."],
        ],
    )

    # 3 Design principles
    _heading(doc, "3. Design Principles")
    _numbered(
        doc,
        [
            "YAML-first — KPI authors declare dimensions, base measures, cuts, and output measures; engine code is frozen.",
            "Catalog freeze — New reusable names go to capabilities/ + registries/; pipeline/ is not edited for new KPIs.",
            "DuckDB for I/O, Pandas for math — Base GROUP BY in DuckDB; all derived/time/cut logic in Pandas.",
            "Anchor-driven time — User-selected month from filters; never max(date) or business_date.",
            "Request-scoped scan — Only requested measures widen lookback; unrequested YAML measures do not expand the scan.",
            "Calendar shifts — 3/6/12-month windows move by calendar months on a dense spine, not by row count.",
            "Explicit measure keys — Every UI column must be declared in measures:; nothing is inferred.",
            "Fail fast — Bind errors at validate time; caps on window rows and trend cells; no silent truncation.",
            "Immutable context — Adapter reads the envelope; calculations never mutate inbound JSON.",
        ],
    )

    # 4 Architecture
    _heading(doc, "4. Architecture Overview")
    _para(doc, "Boundary: the engine is an in-process library, not a standalone service.")
    _table(
        doc,
        ["In scope", "Out of scope (platform)"],
        [
            ["Consume context JSON", "Build context from metadata tables"],
            ["Load KPI/model YAML by kpi_id", "ADLS credentials and path management"],
            ["DuckDB extract (Delta/Parquet)", "Authentication and authorization"],
            ["Pandas calculation + JSON response", "Hierarchy expansion for heir filters"],
            ["validate() dry-run (compile SQL)", "Databricks job orchestration"],
        ],
    )

    _heading(doc, "4.1 Runtime Pipeline", level=2)
    _numbered(
        doc,
        [
            "Adapt context — assert one view; normalize filters; read measure_keys and parameters.",
            "Bind KPI YAML — load kpi_config/kpis/<kpi_id>.yaml; resolve parameters and grain overlay.",
            "Plan time — claim month filter as anchor; compute required_span from requested measures.",
            "DuckDB extract — scan datasets, apply IN filters and date range, GROUP BY to monthly grain.",
            "Pandas — dense month spine, apply cuts, evaluate measure plugins, project requested keys.",
            "Respond — sort, paginate, attach metadata (filters, cuts, trend axes, SQL audit).",
        ],
    )

    _heading(doc, "4.2 Deployment Model", level=2)
    _bullets(
        doc,
        [
            "Entry: kpi_engine.main(context) → compute(context) → JSON.",
            "Host sets udf.module_path to kpi_engine.main for all KPIs.",
            "Copy kpi_engine/ + kpi_config/ into the platform image or volume, or set KPI_ENGINE_CONFIG_DIR.",
            "DuckDB session from platform via HOST_DUCKDB_GETTER; engine never closes host connections.",
            "Kubernetes: scale Deployment replicas; each pod handles independent requests.",
        ],
    )

    # 5 Capabilities
    _heading(doc, "5. Capabilities Catalog")
    _para(
        doc,
        "All callable names are allowlisted in registries/ and documented in CAPABILITIES.md. "
        "Authors pick from this catalog in KPI YAML.",
    )

    for section, names in caps.items():
        if not names:
            continue
        _heading(doc, section, level=2)
        # wrap names in comma-separated groups of ~6 per line as bullets
        chunk = 8
        for i in range(0, len(names), chunk):
            _bullets(doc, [", ".join(f"`{n}`" for n in names[i : i + chunk])])

    _heading(doc, "5.1 Measure Kind Summary", level=2)
    _table(
        doc,
        ["Kind", "Purpose", "Example use"],
        [
            ["point", "Single value at anchor ± offset", "Current month, previous year"],
            ["window", "Trailing / PTD / full-period aggregate", "3-month sum, QTD, YTD"],
            ["trend", "Fixed-length period array for charts", "12-month sparkline"],
            ["arithmetic / fn / expr", "Combine other measures", "YoY %, ratio, blended formula"],
            ["rank / percent_of_total", "Rank or share within a cut", "Category rank, % of total"],
            ["predicate", "1/0 health flag (no row drop)", "Profit OK AND return rate OK"],
            ["hook", "Custom series logic (allowlisted)", "EWMA, seasonal index, streak"],
            ["constant / dimension", "Literal or echo dimension", "Target value, region label"],
        ],
    )

    _heading(doc, "5.2 Row-Level & Entity Features", level=2)
    _bullets(
        doc,
        [
            "Named row steps: expr, lookup, over (lag, running_sum, last_n, row_number, rank).",
            "identity_grain: row helpers as point measures when every emitted cut matches grain.",
            "HAVING: drop groups by measure predicate; optional then_group_by rollup.",
            "Multi-model joins via model_relations in KPI YAML.",
            "SQL models (kind: sql): CTEs and joins shaped in model YAML, not KPI formulas.",
            "Snapshot KPIs: omit time: block for all-history aggregation (no month filter required).",
        ],
    )

    _heading(doc, "5.3 Built-In Aggregations", level=2)
    _para(doc, "base_measures agg: sum, avg, count, min, max, count_distinct, median, percentile, first, last.")

    # 6 Onboarding
    _heading(doc, "6. KPI Onboarding Experience")
    _para(
        doc,
        "Onboarding is designed so data engineers spend time on YAML, not Python. "
        "The standard path requires no engine changes.",
    )

    _heading(doc, "6.1 Standard Path (YAML Only)", level=2)
    _numbered(
        doc,
        [
            "Confirm context: kpi_id, dataset aliases, filter codes, measure_keys, filter_column_mappings.",
            "Model: reuse existing model or add kpi_config/models/<id>.yaml (physical or SQL).",
            "Copy template: cp kpi_config/kpis/sotif/3004.yaml kpi_config/kpis/<kpi_group>/<kpi_id>.yaml.",
            "Fill time, dimensions, base_measures, cuts, measures (every UI measure_key).",
            "Align metadata: measures_required matches YAML; month filter = time.filter_code.",
            "validate(sample_context) — bind + compile SQL without scanning ADLS.",
            "compute(sample_context) — full JSON; add local parquet test under tests/.",
            "Deploy: host module_path = kpi_engine.main (no per-KPI Python file).",
        ],
    )

    _heading(doc, "6.2 Effort Targets", level=2)
    _table(
        doc,
        ["KPI type", "Typical author work", "Engine change"],
        [
            ["Simple SUM/AVG + point/window", "~20 lines YAML, one model", "None"],
            ["Two-model ratio / join", "YAML: two bases + model_relations", "None"],
            ["Complex row pipeline + windows", "YAML: named steps, over, cuts", "None"],
            ["New reusable formula", "capabilities/ + registries/ entry", "Catalog only, not pipeline/"],
            ["Iterative / ML / geospatial", "Hook or SQL model", "Hook or model YAML"],
        ],
    )

    _heading(doc, "6.3 What Authors Never Edit", level=2)
    _bullets(
        doc,
        [
            "kpi_engine/pipeline/ for a standard KPI.",
            "ADLS paths in KPI YAML (paths come from context.datasets).",
            "Per-KPI UDF modules.",
            "DuckDB connection or authentication code.",
        ],
    )

    # 7 JSON contract
    _heading(doc, "7. Response Contract")
    _bullets(
        doc,
        [
            "One row per dimension combination per cut; column per requested measure_key.",
            "Metadata: kpi_id, request_id, anchor, applied/ignored/skipped filters, applied_cuts.",
            "Trend measures return arrays; trend_axes and trend_labels shared in response.",
            "Pagination: total_count, has_more; null page_size returns all rows.",
            "Audit: compiled SQL (parameterized and inlined) in logs per request.",
        ],
    )

    # 8 Scalability
    _heading(doc, "8. Scalability & Operations")
    _table(
        doc,
        ["Dimension", "Approach"],
        [
            ["Concurrent requests", "Stateless compute(context); scale K8s pod replicas"],
            ["Request isolation", "Fresh bind, extract, DataFrames, JSON per call"],
            ["Memory", "Size pods for worst-case extract + pandas (windows, identity grain)"],
            ["I/O bottleneck", "DuckDB scan width = required_span × filters; narrow measures_required"],
            ["Logging", "Per-request log file; KPI_ENGINE_LOG=0 at high volume; use request_id in traces"],
            ["Caching", "Designed (kpi_id + filters + measures + data version); not yet implemented"],
        ],
    )

    # 9 Limitations
    _heading(doc, "9. Limitations & Product Boundaries")
    _para(doc, "These are intentional constraints — not backlog bugs.")

    _heading(doc, "9.1 Identity & Request", level=2)
    _bullets(
        doc,
        [
            "One execution.view_details entry per request.",
            "Empty measures_required computes nothing (does not expand to all YAML measures).",
            "KPI YAML cannot reference another KPI's measures.",
            "context.parameters must match YAML parameters: block or be omitted.",
        ],
    )

    _heading(doc, "9.2 Time & Calendar", level=2)
    _bullets(
        doc,
        [
            "Single anchor month — not a filter range.",
            "business_date on context is ignored for calculations.",
            "No timezone conversion; time.timezone rejected.",
            "Snapshot KPIs (no time:) cannot use windows, trends, or nonzero offsets.",
            "Pick finer than source grain rejected (monthly facts cannot become daily).",
        ],
    )

    _heading(doc, "9.3 Calculation & SQL", level=2)
    _bullets(
        doc,
        [
            "KPI sql:/expr: never enter DuckDB — no SUM(), comments, or subqueries in KPI formulas.",
            "Non-additive aggs (count_distinct, median, percentile) re-read row-level detail.",
            "Trend cells capped at 50,000 per cut; OVER_ROW_CAP = 500,000 detail rows.",
            "heir (hierarchy) filters rejected until expanded upstream.",
            "Regex, JSON, geospatial, ML — hooks or SQL models only.",
        ],
    )

    _heading(doc, "9.4 Extensions", level=2)
    _bullets(
        doc,
        [
            "Custom logic: allowlisted hook names only — no dotted import paths from YAML.",
            "New catalog name: capabilities/ + registries/ — not pipeline/ edits.",
            "New agg, filter operator, or compose template: engine work (rare).",
        ],
    )

    # 10 Security
    _heading(doc, "10. Security")
    _bullets(
        doc,
        [
            "No eval of YAML expressions.",
            "No importlib of caller-supplied module paths.",
            "Parameterized DuckDB queries; filter values never concatenated into SQL.",
            "Dataset paths from trusted context only.",
        ],
    )

    # 11 Rejected ideas
    _heading(doc, "11. Explicitly Rejected Approaches")
    _table(
        doc,
        ["Idea", "Status"],
        [
            ["Pandas eval / df['col'].sum() as measure language", "Rejected"],
            ["DuckDB as loader only (all agg in Pandas)", "Rejected for base agg"],
            ["Anchor = max(time_column) in extract", "Rejected"],
            ["Infer measure_key from catalog op id", "Rejected"],
            ["Standalone API / Databricks job", "Out of scope"],
            ["Replacing YAML with Python classes", "Deferred"],
        ],
    )

    # 12 Glossary
    _heading(doc, "12. Glossary")
    _table(
        doc,
        ["Term", "Meaning"],
        [
            ["Context", "Request JSON from the metadata framework"],
            ["Anchor", "User-selected period; all scalars relative to it"],
            ["required_span", "Date range scanned so lookbacks have history"],
            ["Cut", "Named aggregation grain + filter-ignore policy (e.g. G, R)"],
            ["Measure key", "Column name in JSON; must exist in KPI measures:"],
            ["Capability", "Allowlisted op, function, or hook in registries/"],
            ["Model", "DuckDB extract definition (physical YAML or SQL)"],
        ],
    )

    # 13 References
    _heading(doc, "13. Reference Documentation")
    _bullets(
        doc,
        [
            "kpi-framework-plan.md — locked architecture decisions",
            "kpi-onboarding-guide.md — step-by-step onboarding playbook",
            "kpi-yaml-preparation-guide.md — YAML authoring, limits, when to use what",
            "kpi-yaml-reference.md — every YAML key and op",
            "kpi_engine/registries/CAPABILITIES.md — live capability catalog",
            "README.md — install, run, folder map",
        ],
    )

    _para(doc, f"Document generated {date.today().isoformat()}.", bold=False)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
